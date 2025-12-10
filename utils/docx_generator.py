# utils/docx_generator.py
# DOCX FAYL YARATISH
# python-docx kutubxonasi bilan professional hujjat yaratish

import os
import logging
from typing import Dict, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("⚠️ python-docx kutubxonasi topilmadi!")


class DocxGenerator:
    """
    Professional DOCX hujjatlar yaratish
    Mustaqil ish, referat, kurs ishi uchun
    """

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kutubxonasi o'rnatilmagan. pip install python-docx")

    def create_course_work(
            self,
            content: Dict,
            output_path: str,
            work_type: str = 'mustaqil_ish'
    ) -> bool:
        """
        Mustaqil ish / Referat DOCX yaratish

        Args:
            content: Content dict
            output_path: Chiqish fayl yo'li
            work_type: Ish turi

        Returns:
            bool: Muvaffaqiyat
        """
        try:
            logger.info(f"📄 DOCX yaratish boshlandi: {work_type}")

            # Yangi hujjat
            doc = Document()

            # Stillar sozlash
            self._setup_styles(doc)

            # Titul sahifa
            self._add_title_page(doc, content)

            # Mundarija
            self._add_table_of_contents(doc, content)

            # Kirish
            self._add_introduction(doc, content)

            # Asosiy boblar
            self._add_chapters(doc, content)

            # Xulosa
            self._add_conclusion(doc, content)

            # Tavsiyalar (agar bor bo'lsa)
            if content.get('recommendations'):
                self._add_recommendations(doc, content)

            # Adabiyotlar ro'yxati
            self._add_references(doc, content)

            # Ilovalar (agar bor bo'lsa)
            if content.get('appendix'):
                self._add_appendix(doc, content)

            # Saqlash
            doc.save(output_path)
            logger.info(f"✅ DOCX saqlandi: {output_path}")

            return True

        except Exception as e:
            logger.error(f"❌ DOCX yaratishda xato: {e}")
            return False

    def _setup_styles(self, doc: Document):
        """Hujjat stillarini sozlash"""
        # Normal stil
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        # Paragraf formati
        para_format = style.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_format.first_line_indent = Cm(1.25)
        para_format.space_after = Pt(0)

        # Sarlavha 1
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = 'Times New Roman'
            h1.font.size = Pt(16)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0, 0, 0)
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(12)
            h1.paragraph_format.space_after = Pt(12)

        # Sarlavha 2
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = 'Times New Roman'
            h2.font.size = Pt(14)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0, 0, 0)
            h2.paragraph_format.space_before = Pt(12)
            h2.paragraph_format.space_after = Pt(6)

    def _add_title_page(self, doc: Document, content: Dict):
        """Titul sahifa qo'shish"""
        # O'quv muassasasi
        author_info = content.get('author_info', {})
        institution = author_info.get('institution', "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(institution.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        # Fakultet
        faculty = author_info.get('faculty', '')
        if faculty:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(faculty)
            run.font.size = Pt(14)

        # Kafedra
        department = author_info.get('department', '')
        if department:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(department)
            run.font.size = Pt(14)

        # Bo'sh joy
        for _ in range(4):
            doc.add_paragraph()

        # Sarlavha
        title = content.get('title', 'MAVZU')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = 'Times New Roman'

        # Subtitle
        subtitle = content.get('subtitle', '')
        if subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(14)

        # Bo'sh joy
        for _ in range(8):
            doc.add_paragraph()

        # Yil
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Toshkent – {datetime.now().year}")
        run.font.size = Pt(14)

        # Yangi sahifa
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, content: Dict):
        """Mundarija qo'shish"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MUNDARIJA")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        toc = content.get('table_of_contents', [])
        for item in toc:
            title = item.get('title', '')
            page = item.get('page', '')

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)

            # Tab stop qo'shish
            run = p.add_run(f"{title}")
            run.font.size = Pt(14)

            # Sahifa raqami
            if page:
                # Nuqtalar bilan to'ldirish
                dots = '.' * 50
                p.add_run(f" {dots} {page}")

        doc.add_page_break()

    def _add_introduction(self, doc: Document, content: Dict):
        """Kirish qo'shish"""
        intro = content.get('introduction', {})
        title = intro.get('title', 'KIRISH')
        text = intro.get('content', '')

        # Sarlavha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        doc.add_paragraph()

        # Matn
        paragraphs = text.split('\n\n') if text else []
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                run = p.add_run(para_text.strip())
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'

        doc.add_page_break()

    def _add_chapters(self, doc: Document, content: Dict):
        """Asosiy boblarni qo'shish"""
        chapters = content.get('chapters', [])

        for chapter in chapters:
            number = chapter.get('number', 1)
            title = chapter.get('title', '')

            # Bob sarlavhasi
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{number}-BOB. {title.upper()}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'

            doc.add_paragraph()

            # Bo'limlar
            sections = chapter.get('sections', [])
            for section in sections:
                sec_number = section.get('number', '')
                sec_title = section.get('title', '')
                sec_content = section.get('content', '')

                # Bo'lim sarlavhasi
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(f"{sec_number}. {sec_title}")
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'

                # Bo'lim matni
                paragraphs = sec_content.split('\n\n') if sec_content else [sec_content]
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph()
                        p.paragraph_format.first_line_indent = Cm(1.25)
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                        run = p.add_run(para_text.strip())
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'

                doc.add_paragraph()

            doc.add_page_break()

    def _add_conclusion(self, doc: Document, content: Dict):
        """Xulosa qo'shish"""
        conclusion = content.get('conclusion', {})
        title = conclusion.get('title', 'XULOSA')
        text = conclusion.get('content', '')

        # Sarlavha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        doc.add_paragraph()

        # Matn
        paragraphs = text.split('\n\n') if text else []
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                run = p.add_run(para_text.strip())
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'

        doc.add_page_break()

    def _add_recommendations(self, doc: Document, content: Dict):
        """Tavsiyalar qo'shish"""
        recommendations = content.get('recommendations', [])

        if not recommendations:
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TAVSIYALAR")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            run = p.add_run(f"{i}. {rec}")
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        doc.add_page_break()

    def _add_references(self, doc: Document, content: Dict):
        """Adabiyotlar ro'yxati"""
        references = content.get('references', [])

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FOYDALANILGAN ADABIYOTLAR RO'YXATI")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        for ref in references:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)

            run = p.add_run(ref)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

    def _add_appendix(self, doc: Document, content: Dict):
        """Ilovalar qo'shish"""
        appendix = content.get('appendix')

        if not appendix:
            return

        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ILOVALAR")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        if isinstance(appendix, str):
            p = doc.add_paragraph()
            run = p.add_run(appendix)
            run.font.size = Pt(14)
        elif isinstance(appendix, list):
            for item in appendix:
                p = doc.add_paragraph()
                run = p.add_run(str(item))
                run.font.size = Pt(14)


# Helper function
def create_docx_from_content(content: Dict, output_path: str, work_type: str = 'mustaqil_ish') -> bool:
    """DOCX yaratish helper"""
    try:
        generator = DocxGenerator()
        return generator.create_course_work(content, output_path, work_type)
    except Exception as e:
        logger.error(f"❌ DOCX helper xato: {e}")
        return False
